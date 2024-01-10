import csv
import os
import re

import pandas as pd
from docx import Document
from py2neo import Graph, Node, Relationship
from tqdm import tqdm

from handleColonGpt import process_text_ignore_extra_content


# word 要求
# word的子标题下面 如果是正文，则必须只有一段
# 在一段正文中，如果有编号“（一）” ，这种，则必须只有一个编号，不可以同时存在多个并列关系，说人话就是不能有两个或者多个（一）  “设备维护主要包括：（一）表单建立。（二）指派人员。 xxx 网络信息管理员的目的：（一）对（二）综合”
# 正文里面有冒号 当且仅当有（一） 这类标号
# 一级标题 第xx章 为行动大类  类似于 第五章 进攻行动
# 二级标题 第xx节 中为具体行动的名称，类似于 第二节  进攻掩护行动
# 三级标题【】中为具体行动的特征/要素，类似于 【对友方效果】，且 不允许将正文的内容放于三级标题格式
# 正文中 如果有（一）这类标号则必须使用。（句号）进行结尾。类似于“设备维护主要包括：（一）表单建立。（二）指派人员。”，而不能是；（分号）类似于“类似于“设备维护主要包括：（一）表单建立；（二）指派人员。”
# word中比如使用连续的子标题，即一级标题下面使用二级标题，二级标题下面使用三级标题，而不能 一级标题下面直接跟三级标题

class readWord:
    def __init__(self, word_path, csv_file_path):
        self.titleDict = {
            '1级': 'Heading 1',
            '2级': 'Heading 2',
            '3级': 'Heading 3',
            '4级': 'Heading 4',
            '5级': 'Heading 5',
            '6级': 'Heading 6',
            '7级': 'Heading 7',
            '8级': 'Heading 8',
            '9级': 'Heading 9',
            '10级': 'Heading 10',
        }
        self.doc = Document(word_path)

        self.all_tables = self.read_all_tables_to_dataframe()
        '''for table in self.all_tables:
            print(table)'''
        # print(len(self.all_tables))
        self.tablesindex = 0
        self.visit = [False] * len(self.doc.paragraphs)
        self.csv_file_path = csv_file_path
        if not os.path.exists(csv_file_path):
            print(f"{csv_file_path} not exist! created!")

        """""with open(self.csv_file, 'w', newline='') as file:
            # 创建 CSV writer 对象
            self.csv_writer = csv.writer(file)"""""
        self.rows = []

    def print_content(self):
        # 获取word内容以及 heading
        for pa in self.doc.paragraphs:
            print(pa.style.name)
            print(pa.text)

    def read_all_tables_to_dataframe(self):
        # 初始化一个空列表来存储所有表格的DataFrame
        all_tables = []
        # 遍历文档中的所有表格
        for table in self.doc.tables:
            # 将表格转换为DataFrame
            table_df = self.table_to_dataframe(table)
            all_tables.append(table_df)

        return all_tables

    def table_to_dataframe(self, table):
        # 初始化列名为空列表
        columns = []

        # 初始化一个空列表用于存储行数据
        rows_data = []

        # 遍历表格中的行
        for i, row in enumerate(table.rows):
            # 如果是第一行，将其作为列名
            if i == 0:
                columns = [cell.text.strip() for cell in row.cells]
            else:
                # 否则，将行数据添加到列表
                row_data = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_data)

        # 一次性创建DataFrame
        return pd.DataFrame(rows_data, columns=columns)

    def read_first_level_titles(self):
        # 初始化标题列表
        first_level_titles = []
        first_level_titles_index = []

        for i in range(len(self.doc.paragraphs)):
            # print(self.doc.paragraphs[i].style.name)
            # print('$' * 100)
            curParagraph = self.doc.paragraphs[i]
            if curParagraph.style.name.startswith('Heading 1') or curParagraph.style.name == '1级':
                # 获取一级标题文本
                # print("get one title")
                title_text = self.extract_content_inside_brackets(curParagraph.text) or curParagraph.text
                first_level_titles.append(title_text)
                first_level_titles_index.append(i)

        return first_level_titles, first_level_titles_index

    def read_titles(self):
        all_titles_dict = {}
        first_level_titles, first_level_titles_index = self.read_first_level_titles()
        # print(first_level_titles)
        for first_title, i in zip(first_level_titles, first_level_titles_index):
            all_titles_dict[first_title] = self.read_all_titles(current_level=1, start=i)

        # return all_titles_dict
        """dic = {}
        all_titles_dict[first_title] = self.read_all_titles2(plevel=0, pdict=dic, start=0)"""
        return all_titles_dict

    def read_all_titles(self, current_level, start, fillTable=False):
        current_dict = {}
        for i in range(start + 1, len(self.doc.paragraphs)):
            curParagraph = self.doc.paragraphs[i]  # 当前段落
            style_name = curParagraph.style.name  # 当前段落格式
            if style_name in self.titleDict.keys():
                style_name = self.titleDict[style_name]
            # print(style_name)
            # 如果需要进行填表，且当前表没有被填过，或者当前段落不是标题（标题 x或者 x级）
            if not self.visit[i] and style_name == "Normal":
                self.visit[i] = True
                return process_text_ignore_extra_content(
                    curParagraph.text) if '（一）' in curParagraph.text else [curParagraph.text]
            if self.visit[i] or not style_name.startswith('Heading'):
                # 跳过非标题段落
                continue
            # print(style_name.split())
            level = int(style_name.split()[-1])

            if level > current_level:
                self.visit[i] = True
                title_text = self.extract_content_inside_brackets(curParagraph.text) or curParagraph.text
                current_dict[title_text] = self.read_all_titles(level, i)
            else:
                break
        if fillTable and len(current_dict) == 0:
            current_dict = self.all_tables[self.tablesindex]
            self.tablesindex += 1
        return current_dict

    def extract_content_inside_brackets(self, sentence):
        """
        提取中文方括号内的内容

        Parameters:
        sentence (str): 包含方括号的句子

        Returns:
        str: 方括号内的内容，如果找不到则返回 None
        """
        p = re.compile(r'[\[({【](.*?)[\])}】]', re.S)
        matches = re.findall(p, sentence)
        return matches[0] if len(matches) > 0 else None

    def format_data(self, data, indent=4):
        result = ""
        for key, value in data.items():
            result += " " * indent + f"{key}:\n"
            if isinstance(value, dict):
                result += self.format_data(value, indent + 4)
            else:
                result += " " * (indent + 4) + str(value) + "\n"
        return result

    def loadCsv(self, data, parent=None):
        # 写入数据

        for key, value in data.items():
            # node = Node("Chapter", name=key)
            # graph.create(node)
            if parent is not None:
                parentType = "Chapter"
                keyType = "Chapter"
                if parent == '全文':
                    parentType = "All"
                elif '章' in parent and '第' in parent:
                    parentType = "Chapter"
                elif '节' in parent and '第' in parent:
                    parentType = "movementType"
                else:
                    parentType = "movement"

                if '章' in key:
                    keyType = "Chapter"
                elif '节' in key and '第' in key:
                    keyType = "movementType"
                else:
                    keyType = "movement"
                self.rows.append([parent, "HAS_CHILD", key, '_', '_', parentType, keyType])
                # csv_writer.writerow((parent, "HAS_CHILD", key, '_', '_', "Chapter", "Chapter"))
                # relationship = Relationship(parent, "HAS_CHILD", node)
                # graph.create(relationship)
            if isinstance(value, dict):
                self.loadCsv(value, key)
            elif isinstance(value, list):
                for item in value:
                    self.rows.append([key, "HAS_CHILD", item, '_', '_', "Chapter", "Content"])
            else:
                self.rows.append([key, "HAS_CHILD", value, '_', '_', "Chapter", "Content"])

    def getTitleCsv(self, data):
        # graph.delete_all()
        paper = {}
        paper['全文'] = data
        print(paper)

        with open(self.csv_file_path, 'w', encoding='utf-8') as file:
            self.loadCsv(data=paper)
        self.assign_numbers()
        self.create_graph_from_csv(self.csv_file_path)

    def create_graph_from_csv(self, csv_file: str, uri: str = "bolt://localhost:7687", username: str = "neo4j",
                              password: str = "12345678", name: str = "neo4j",
                              deleteAll: bool = True):
        '''

        :param csv_file:  需要导入的csv
        :param uri: 图数据库地址
        :param username: 用户名
        :param password: 密码
        :param name: 数据库名称
        :param deleteAll: 是否在导入以前清除所有点 默认为True
        :return:
        '''
        graph = Graph(uri, auth=(username, password), name=name)
        if deleteAll:
            graph.delete_all()  # 清除neo4j中原有的结点等所有信息
        created_nodes = {}

        with open(csv_file, 'r', encoding='utf-8') as file:
            csv_reader = csv.reader(file)
            # 跳过表头
            # next(csv_reader, None)

            for row in tqdm(csv_reader, position=0):
                print(row)
                sender_name, action_name, receiver_name, sender_number, receiver_number, sender_class, receiver_class = row

                # 检查节点是否已存在，如果不存在则创建
                sender_node = created_nodes.get(sender_number,
                                                Node(sender_class, name=sender_name, number=sender_number))
                receiver_node = created_nodes.get(receiver_number,
                                                  Node(receiver_class, name=receiver_name, number=receiver_number))

                # 将节点添加到已创建的节点字典
                created_nodes[sender_number] = sender_node
                created_nodes[receiver_number] = receiver_node

                # 创建关系
                action_relationship = Relationship(sender_node, action_name, receiver_node)

                # 将节点和关系添加到图数据库
                graph.create(sender_node | receiver_node | action_relationship)

    def assign_numbers(self, current_number=0):
        '''

        Args:
            input_file:
            output_file:
            current_number: 当前并列连词编号

        Returns:

        '''
        objects = {}  # 用于存储实物及其编号的字典
        # current_number  # 当前编号553

        # print(reader)
        with open(self.csv_file_path, 'w', newline='', encoding='utf-8') as outfile:
            writer = csv.writer(outfile)

            for row in tqdm(self.rows, position=0):
                # 获取实物名称
                # 检查行的长度，如果不足5个字段则跳过

                """if len(row) < 5:
                    continue"""
                # print(row)
                # 获取实物名称
                print(row)
                object1, relation, object2, number1, number2 = row[:5]

                # 处理第一列实物
                if object1 not in objects:
                    if number1 == '_':
                        objects[object1] = current_number
                        current_number += 1
                    else:
                        objects[object1] = number1

                # 处理第二列实物
                if object2 not in objects:
                    if number2 == '_':
                        objects[object2] = current_number
                        current_number += 1
                    else:
                        objects[object2] = number2

                # 更新行中的编号
                row[3] = objects[object1] if number1 == '_' else number1
                row[4] = objects[object2] if number2 == '_' else number2
                writer.writerow(row)
                # 写入到输出文件
        # print(rows)


if __name__ == "__main__":
    # 生成样本文档
    word_path = "格式模板样例-tl（公开）——改标题三.docx"
    # create_sample_document(word_path)
    csv_file = 'csv_file.csv'
    rd = readWord(word_path, csv_file)
    res = rd.read_titles()
    print(rd.format_data(res))
    rd.getTitleCsv(res)
    print(rd.format_data(data=res))

    # rd.print_content()
