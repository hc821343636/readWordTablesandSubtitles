from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# 创建一个新的Word文档
doc = Document()

# 设置文档标题
title = doc.add_heading(level=0)
title_run = title.add_run('容器安全报告')
title_run.font.size = Pt(16)
title_run.font.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 添加报告概览部分
overview = doc.add_heading(level=1)
overview_run = overview.add_run('报告概览')
overview_run.font.size = Pt(14)
overview_run.font.bold = True

# 报告概览的具体内容
details = [
    "报告日期：2023年12月22日",
    "报告版本：1.0",
    "扫描对象：容器云平台",
    "扫描范围：全部容器镜像及运行时环境"
]
for detail in details:
    para = doc.add_paragraph()
    para_run = para.add_run(detail)
    para_run.font.size = Pt(12)

# 添加漏洞摘要部分
summary = doc.add_heading(level=1)
summary_run = summary.add_run('漏洞摘要')
summary_run.font.size = Pt(14)
summary_run.font.bold = True

# 漏洞摘要的具体内容
vulnerabilities = [
    "严重漏洞：4个",
    "高危漏洞：10个",
    "中危漏洞：15个",
    "低危漏洞：20个"
]
for vuln in vulnerabilities:
    para = doc.add_paragraph()
    para_run = para.add_run(vuln)
    para_run.font.size = Pt(12)

# 添加重点漏洞详情部分
details = doc.add_heading(level=1)
details_run = details.add_run('重点漏洞详情及建议措施')
details_run.font.size = Pt(14)
details_run.font.bold = True

# 重点漏洞的具体内容
vuln_details = [
    "CVE-2023-0001: SQL注入\n  严重性：严重\n  影响组件：数据库服务\n  详细建议：更新至最新版本的数据库软件以修复已知漏洞。实施严格的输入验证措施。使用参数化查询和预处理语句。对数据库访问进行最小权限配置。\n  状态：待修复",
    "CVE-2023-0025: 跨站脚本攻击 (XSS)\n  严重性：高\n  影响组件：Web应用界面\n  详细建议：对所有用户输入进行字符过滤和转义处理。实施内容安全策略（CSP）。使用安全工具进行代码审查。提高开发人员的安全编码意识。\n  状态：修复中",
    "CVE-2023-0018: 容器逃逸\n  严重性：严重\n  影响组件：容器运行时\n  详细建议：更新容器管理工具和运行时环境。对容器进行资源限制和隔离。实施角色基于访问控制（RBAC）。定期检查和更新容器的安全配置。\n  状态：待修复"
]
for vuln_detail in vuln_details:
    para = doc.add_paragraph()
    para_run = para.add_run(vuln_detail)
    para_run.font.size = Pt(12)

# 保存文档
doc.save('Container_Security_Report.docx')
